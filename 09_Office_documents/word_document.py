import docx

docum = docx.Document('demo.docx')
print(docum.paragraphs[0])
print(docum.paragraphs[0].text)

par = docum.paragraphs[1]
print(par.runs)
print(par.runs[0].text)

print(par.runs[1].bold)
print(par.runs[3].italic)

par.runs[0].underline = True
par.runs[2].text = ' italic and underlined '
docum.save('demo2.docx')

print(par.style)
par.style = 'Title'
docum.save('demo3.docx')

# ============================================

doc = docx.Document()
doc.add_paragraph('Hello, this is a paragraph. ')
doc.add_paragraph('This is another paragraph. ')
par = doc.paragraphs[0]
par.add_run('This is a new run and it will be bold. ')
par.runs[1].bold = True
doc.save('demo4.docx')
